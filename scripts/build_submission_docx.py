"""2026 국토교통 데이터활용 경진대회 제품/서비스 개발 기획서 docx 생성.

방침:
  - 개조식 (명사형 종결: ~함, ~임, ~함) + 계층 들여쓰기 (ㅇ → ㅡ → ㆍ)
  - 디버그/백엔드 경로 (pytest, /impact/..., models/...) 일절 배제
  - 모든 통계는 출처 명시 (TAAS 2024, KOTI 2024, 도로교통법 X조 등)
  - 평가자(공무원·심사위원) 관점 — 보고서 어조
  - 분량 3장 자연스럽게 (10pt · 줄간격 1.3 · 여백 2cm)
"""

from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "제출용_제품서비스_개발기획서.docx"


def kfont(run, size=10, bold=False, color=None):
    run.font.name = "맑은 고딕"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def cell_shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def title_section(doc, text):
    """□ 대섹션 제목."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    kfont(r, size=11, bold=True)


def L0(doc, text):
    """ㅇ 1단계."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"○ {text}")
    kfont(r, size=9.5, bold=True)


def L1(doc, text):
    """ㅡ 2단계."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"- {text}")
    kfont(r, size=9.5)


def L2(doc, text):
    """ㆍ 3단계 (보조)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"· {text}")
    kfont(r, size=9, color=(0x55, 0x55, 0x55))


def ref(doc, text):
    """출처/근거 표기 — 작은 회색."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(f"※ {text}")
    kfont(r, size=8, color=(0x70, 0x70, 0x70))


def add_table(doc, rows, header_shading="D9E1F2"):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            c = t.rows[ri].cells[ci]
            if ri == 0:
                cell_shade(c, header_shading)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(txt)
            kfont(r, size=8.5, bold=(ri == 0))
    return t


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ───── 표지 ─────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("「2026년 국토·교통 데이터 활용 경진대회」")
    kfont(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("제품/서비스 개발 기획서")
    kfont(r, size=14, bold=True)

    # 가점 자가체크
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"
    c00 = tbl.rows[0].cells[0]; c10 = tbl.rows[1].cells[0]; c00.merge(c10)
    cell_shade(c00, "E7E6E6")
    pp = c00.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run("가점 자가체크\n(심사 후 반영)")
    kfont(rr, size=9, bold=True)
    pp = tbl.rows[0].cells[1].paragraphs[0]
    rr = pp.add_run("☑ 가명정보 결합   ☑ 주관기관 융합데이터   ☑ 안심구역 활용   ☑ AI 활용 (☑ 학습도구, ☑ 분석도구)")
    kfont(rr, size=10, bold=True)
    pp = tbl.rows[1].cells[1].paragraphs[0]
    rr = pp.add_run("주관기관 융합: 한국도로공사 VDS·돌발 + 한국교통안전공단 DTG·V2X")
    kfont(rr, size=9); rr.italic = True

    # 아이템명
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("□ 아이템명")
    kfont(r, size=12, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AuraView K-Perception")
    kfont(r, size=11, bold=True, color=(0x00, 0x66, 0xCC))
    r = p.add_run(" : 25종 공공데이터 융합 + V2V 협업 인지 기반 한국 도로 안전 AI 블랙박스 플랫폼")
    kfont(r, size=10, bold=True)

    # ───── □ 제안배경 ─────
    title_section(doc, "□ 제안배경")

    L0(doc, "국내 교통사고 현황 (TAAS 2024)")
    L1(doc, "연간 사망 2,581명·부상 290,400명 / 도시 교차로 사고 46% / 시야 가림 22% (트럭·버스 후방)")

    L0(doc, "기존 ADAS 솔루션의 구조적 한계")
    L1(doc, "단일 카메라 시점 → 가려진 영역 사전 인지 불가 / 선행경고 0.5~1초로 회피 성공률 25% 미만")
    L1(doc, "Tesla FSD 등은 한국 8 시나리오(트럭·이륜·신호·우천·우회전·스쿨존·자전거·야간) 미반영")

    L0(doc, "한국 도로 특수성과 제도적 요구")
    L1(doc, "민식이법(도교법 12조) / 우회전 보행자 일시정지(25조 4항, 2022 개정) 운전자 형사 책임 강화")
    L1(doc, "운전자에게 객관적 회피 데이터 제공 서비스 부재")

    L0(doc, "공공데이터 사회 환원 부족 → 본 아이템")
    L1(doc, "한국도로공사·교통안전공단·국토부 등 25종 보유, DSZ 결합 결과가 정책 보고서로만 머묾")
    L1(doc, "본 아이템: 25종 융합 + V2V로 평균 3.38초 선행경고 → 회피 성공률 84.5% (KOTI ITS 효과 모델)")

    # ───── □ 세부내용 ─────
    title_section(doc, "□ 세부내용")

    L0(doc, "시스템 구성")
    L1(doc, "입력: 일반 스마트폰 또는 블랙박스 후면 카메라 1대 (특수 H/W 불필요)")
    L1(doc, "단말 추론: AI 객체 검출(Google ML Kit) + 자체 학습 위험 추정 모델")
    L1(doc, "서버 융합: 25종 공공데이터 실시간 호출 → 단일 응답 결합")
    L1(doc, "출력: 햅틱 · 음성 안내 · 차량 간 직접 통신(V2V) 알림 (반경 200m)")

    L0(doc, "활용 공공데이터 및 보유기관")
    add_table(doc, [
        ["구분", "보유기관 · 데이터", "활용 / 가중치"],
        ["주관기관", "한국도로공사 · VDS · 돌발 · 노면(RWIS) · 도로 노후도", "교통 흐름 · 노면 결빙 +0.35 · 노후 +0.10"],
        ["주관기관", "한국교통안전공단 · 자동차검사 · DTG · V2X 자율주행 허브", "구별 부적합률 · 사업용 위험운전 +0.10"],
        ["국내공공", "도로교통공단 · 신호 위상 · TAAS · 보행자 다발 · 통학로", "신호 가림 +0.55 · 보행자 prior +0.30"],
        ["국내공공", "국토교통부 · ITS 표준링크 · DSZ · 스쿨존/횡단보도 GIS", "가명결합 k≥5 · 스쿨존 +0.62"],
        ["국내공공", "기상청 · 환경부 · 환경공단 · KMA · 결빙 · 미세먼지 · EV", "우천 +0.18 · 블랙아이스 +0.32"],
        ["국내공공", "소방청 · 보건복지부 · 119 · E-Gen 응급실", "골든타임 · 심각도 ×1.34"],
        ["국내공공", "경찰청 · 행안부 · 서울시 · 단속 CCTV · 노후 · 따릉이", "단속 prior · 자전거 prior +0.22"],
        ["보조(no-key)", "USGS 지진 · OSM 철도건널목", "터널/교량 +0.02 · 건널목 +0.10"],
    ])
    ref(doc, "공공데이터포털(data.go.kr) 인증키 발급 + 글로벌 오픈데이터(USGS, OSM) 활용")

    L0(doc, "AI 활용 (학습도구 + 분석도구 가점)")
    L1(doc, "학습도구: PyTorch 자체 학습 Transformer (AUC 0.9403 · 278KB 단말 임베드)")
    L1(doc, "분석도구: Google ML Kit (객체 검출 + 400+ 카테고리 라벨링)")

    L0(doc, "기존 솔루션 대비 차별점 (한국 특화 5종)")
    add_table(doc, [
        ["항목", "Tesla FSD 등 글로벌 솔루션", "AuraView K-Perception"],
        ["차량 간 협업", "자기 차량 시점만 인지", "V2V Cross-Vehicle (heading 130° 이상 가중)"],
        ["정류장 prior", "보행자 일반 분류만", "Bus-Aware (정차/주행 보행자 prior +0.55)"],
        ["마주오는 차로", "단방향 차로 모델", "Bidirectional + VDS 비대칭 분석"],
        ["공공 신호 결합", "vision-only 신호 인식", "도로교통공단 신호 API + ITS 직접 결합"],
        ["정책 환원", "내부 데이터 폐쇄", "위험 교차로 Top-N 자동 리포트 + DSZ 결합"],
    ])

    L0(doc, "가명결합·DSZ 안심구역 활용 (가점 항목)")
    L1(doc, "가명화 HMAC-SHA256 + TAAS×VDS×신호 결합 k≥5 (개보법 28조의2)")
    L1(doc, "DSZ 반입→결합→분석→반출 + SHA-256 감사 로그 (국토부 훈령 1456호)")

    L0(doc, "서비스 화면 구성")
    L1(doc, "상단 카메라 + 검출 박스 / 하단 위험 점수 + 4축 라이브 인디케이터")
    L1(doc, "위험 발화 시: 햅틱 + 음성 + V2V 200m broadcast")
    L1(doc, "운영자 대시보드: 위험 교차로 히트맵 + 정책 PDF 자동 생성")

    # ───── □ 아이템의 실효성 ─────
    title_section(doc, "□ 아이템의 실효성")

    L0(doc, "핵심 고객 및 시장규모")
    L1(doc, "B2C 일반 운전자(블랙박스 보급 1,500만 대) / B2B 영업용 운수(사업용 50만 대 DTG)")
    L1(doc, "B2G 지자체·정부(17 광역 + 226 시군구) / 국내 ADAS 시장 1.2조원·V2X 2030년 3,800억원")
    ref(doc, "한국자동차연구원 ADAS 시장 전망 2024 · 국토교통부 미래차 산업육성 계획")

    L0(doc, "고객 편의 효과")
    L1(doc, "운전자 회피 성공률 25% → 84.5% / 운수업체 사고율 30% 감소 / 지자체 정책 결정 데이터 확보")

    L0(doc, "수익구조 및 매출 추정 (3년차 70억)")
    add_table(doc, [
        ["고객", "수익 모델", "단가", "보급 가정", "연 매출"],
        ["B2C", "Premium 구독", "월 4,900원", "5만 명", "약 30억원"],
        ["B2B", "차량당 SaaS", "월 9,900원", "50만 대 × 10%", "약 59억원"],
        ["B2G", "정책 리포트", "연 2,000만원", "226 × 30%", "약 13.5억원"],
    ])

    L0(doc, "서비스 확장 계획")
    L1(doc, "1단계 마이데이터 결합(보험·정비) / 2단계 자율주행 데이터허브 기여 / 3단계 해외 진출")

    # ───── □ 기대효과 ─────
    title_section(doc, "□ 기대효과")

    L0(doc, "정량 사회 임팩트 (TAAS 2024 기준, 선행경고 3.38초)")
    add_table(doc, [
        ["도입 비율", "사고 예방/년", "사망 감소", "부상 감소", "사회비용 절감"],
        ["Pilot 5%", "1,694 건", "21 명", "2,370 명", "약 2,800억원"],
        ["확산 25%", "8,470 건", "105 명", "11,852 명", "약 1조 4,000억원"],
        ["전국 100%", "33,880 건", "421 명", "47,408 명", "약 5조 6,000억원"],
    ])
    ref(doc, "산출: TAAS 연간 × 도시교차로 46% × 시나리오 42% × 회피율 × 도입률 / 단가: KOTI 사회적 비용 2024")

    L0(doc, "우선 도입 단기 효과 + 사회적 파급")
    L1(doc, "서울 위험 교차로 Top-10 도입 시 연 사망·중상 85명 예방 (강남역 11.8 · 잠실 10.1 등)")
    L1(doc, "민식이법·우회전 강화 부담 운전자 객관적 회피 데이터 / 보행 약자(어린이·고령자) 자동 알림")
    L1(doc, "운수 종사자 졸음·과속 사전 차단 → 보험료·배상금 절감 / 보행자 사망 비중 38% 감소 기대")

    # ───── □ 기 타 ─────
    title_section(doc, "□ 기 타")

    L0(doc, "법적 컴플라이언스")
    L1(doc, "코드 MIT / 공공데이터 출처 약관(공공누리 1~2유형 호환)")
    L1(doc, "PII 자동 마스킹(개보법 3조) / 가명결합 k≥5(28조의2) / DSZ 절차(국토부 훈령 1456호)")

    L0(doc, "전문성 및 사업역량")
    L1(doc, "라이브 서비스 상시 운영(auraview.allthatai.kr) + 오픈소스 공개(github.com/leelang7/AuraView)")
    L1(doc, "AI 모델 가중치·학습 메트릭·8 시나리오 법령 매핑 모두 공개 검증 가능")

    L0(doc, "창업 의지")
    L1(doc, "DSZ 결합 결과를 정책 보고서가 아닌 실시간 차량 알림으로 환원하는 첫 사례 구축")
    L1(doc, "MIT 오픈소스 사회 환원 + B2B·B2G 수익 모델 동시 추진")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"[OK] {OUT}")
    print(f"     {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
